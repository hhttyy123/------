"""将 docs/使用手册.md 转换为 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm
import re

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)

with open(r'D:\人力资源项目\docs\使用手册.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()


def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'


def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    run.bold = bold


def add_table_simple(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(9)
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    doc.add_paragraph()


def clean_line(text):
    """移除 Markdown 格式标记"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if not line:
        i += 1
        continue

    # H1
    if line.startswith('# ') and not line.startswith('## '):
        add_heading_styled(line[2:], 0)
        i += 1
        continue

    # H2
    if line.startswith('## '):
        add_heading_styled(line[3:], 1)
        i += 1
        continue

    # H3
    if line.startswith('### '):
        add_heading_styled(line[4:], 2)
        i += 1
        continue

    # Markdown table
    if line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        if len(table_lines) >= 2:
            headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            rows_data = [
                [c.strip() for c in row.split('|')[1:-1]]
                for row in table_lines[2:]
            ]
            add_table_simple(headers, rows_data)
        continue

    # Blockquote
    if line.startswith('> '):
        add_para(clean_line(line[2:]), bold=True, indent=True)
        i += 1
        continue

    # Bullet
    if line.startswith('- '):
        add_para('• ' + clean_line(line[2:]), indent=True)
        i += 1
        continue

    # Numbered
    if re.match(r'^\d+\.\s', line):
        add_para(clean_line(line))
        i += 1
        continue

    # Bold-only line
    if line.startswith('**'):
        add_para(clean_line(line), bold=True)
        i += 1
        continue

    # Normal paragraph
    add_para(clean_line(line))
    i += 1

output_path = r'D:\人力资源项目\docs\使用手册.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
